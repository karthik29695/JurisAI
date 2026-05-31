"""
app/services/ocr_service.py  —  Text extraction from PDFs, images, and DOCX files.

Extraction chain (per format)
──────────────────────────────
PDF   → PyMuPDF text layer  →  Tesseract OCR  →  Gemini Vision OCR
Image → Tesseract OCR       →  Gemini Vision OCR
DOCX  → python-docx paragraphs + table cells
"""
import io
import logging
from typing import Optional

log = logging.getLogger("jurisai.ocr")

# ── Optional imports ──────────────────────────────────────────────────────────
try:
    import fitz          # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    fitz = None
    FITZ_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    Image = None
    PIL_AVAILABLE = False

try:
    import pytesseract
    HAVE_TESS = True
except ImportError:
    pytesseract = None
    HAVE_TESS = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    DOCX_AVAILABLE = False

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".tiff", ".tif", ".bmp"}
DOCX_EXTS  = {".doc", ".docx"}


class OCRService:
    """
    Extracts plain text from uploaded legal documents.
    Inject a GeminiService to enable Gemini Vision OCR fallback.
    """

    def __init__(self, gemini_service=None):
        self._gemini = gemini_service

    # ── Gemini Vision OCR ─────────────────────────────────────────────────────

    def _ocr_with_gemini(self, pil_image) -> str:
        if not self._gemini or not self._gemini.available:
            return ""
        try:
            resp = self._gemini.generate(
                ["Extract all text from this image. Return only the plain text, no commentary.",
                 pil_image]
            )
            return (resp.text or "").strip()
        except Exception as exc:
            log.warning("Gemini OCR failed: %s", exc)
            return ""

    # ── PDF ───────────────────────────────────────────────────────────────────

    def extract_from_pdf(self, file_bytes: bytes) -> str:
        """Extract text from a PDF using PyMuPDF + OCR fallback."""
        if not FITZ_AVAILABLE:
            raise RuntimeError("PyMuPDF not installed. Run: pip install pymupdf")
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        parts = []
        for page in doc:
            txt = (page.get_text("text") or "").strip()
            if len(txt) >= 40:
                parts.append(txt)
                continue
            # Render page to image for OCR
            try:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img = Image.open(io.BytesIO(pix.tobytes("png")))
            except Exception:
                parts.append(txt)
                continue
            ocr = ""
            if HAVE_TESS:
                try:
                    ocr = pytesseract.image_to_string(img, lang="eng")
                except Exception:
                    pass
            if not ocr.strip():
                ocr = self._ocr_with_gemini(img)
            parts.append((ocr or "").strip())
        doc.close()
        return "\n\n".join(p for p in parts if p)

    # ── Image ─────────────────────────────────────────────────────────────────

    def extract_from_image(self, file_bytes: bytes, filename: str = "") -> str:
        """Extract text from an image via Tesseract → Gemini Vision chain."""
        if not PIL_AVAILABLE:
            raise RuntimeError("Pillow not installed. Run: pip install Pillow")
        img = Image.open(io.BytesIO(file_bytes))
        if HAVE_TESS:
            try:
                txt = pytesseract.image_to_string(img, lang="eng")
                if txt.strip():
                    return txt.strip()
            except Exception as exc:
                log.warning("Tesseract failed for '%s': %s", filename, exc)
        txt = self._ocr_with_gemini(img)
        if txt.strip():
            return txt.strip()
        raise ValueError(f"No text extracted from image '{filename}'.")

    # ── DOCX ──────────────────────────────────────────────────────────────────

    def extract_from_docx(self, file_bytes: bytes) -> str:
        """Extract text from a Word document using python-docx."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
        doc = DocxDocument(io.BytesIO(file_bytes))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paras.append(cell.text.strip())
        return "\n\n".join(paras)

    # ── Dispatch ──────────────────────────────────────────────────────────────

    def extract(self, file_bytes: bytes, filename: str, ext: str) -> str:
        """Route extraction based on file extension."""
        ext = ext.lower()
        if ext == ".pdf":
            return self.extract_from_pdf(file_bytes)
        if ext in IMAGE_EXTS:
            return self.extract_from_image(file_bytes, filename)
        if ext in DOCX_EXTS:
            return self.extract_from_docx(file_bytes)
        raise ValueError(f"Unsupported file type: '{ext}'")
